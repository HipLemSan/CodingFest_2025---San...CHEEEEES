import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import sqlite3
from datetime import datetime
import pandas as pd
from docx import Document
import os
import shutil

class FilamentApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Учет материалов для 3D принтеров")
        self.root.geometry("1300x750")
        self.root.configure(bg='white')
        
        self.current_db = "filament.db"
        self.conn = None
        self.cursor = None
        
        self.material_filter = None
        self.diameter_filter = None
        self.color_filter = None
        self.status_filter = None
        self.search_var = None
        self.tree = None
        
        self.setup_database()
        self.create_interface()
        self.load_data()

    def setup_database(self):
        try:
            self.conn = sqlite3.connect(self.current_db)
            self.cursor = self.conn.cursor()
            
            self.cursor.execute('''
                CREATE TABLE IF NOT EXISTS materials (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    date TEXT NOT NULL,
                    material_type TEXT NOT NULL,
                    spool_size TEXT NOT NULL,
                    diameter TEXT NOT NULL,
                    color TEXT NOT NULL,
                    storage_conditions TEXT NOT NULL,
                    status TEXT NOT NULL,
                    remaining_weight REAL NOT NULL,
                    notes TEXT
                )
            ''')
            self.conn.commit()
            
            self.cursor.execute("SELECT COUNT(*) FROM materials")
            if self.cursor.fetchone()[0] == 0:
                self.add_sample_data()
                
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка базы: {str(e)}")

    def add_sample_data(self):
        sample_data = [
            ('17.11.2024', 'ABS', '750 г (1 кг), катушка Ø 100 мм', '1.75 мм', 'чёрный', 
             'В герметичном контейнере, с силикагелем, при t 15–25°C, влажность ≤ 40%', 
             'Добавлен', 1.0, ''),
            ('17.11.2024', 'PLA', '750 г (1 кг), катушка Ø 100 мм', '1.75 мм', 'белый', 
             'В сухом, прохладном месте, избегать прямого света, t 15–25°C, влажность ≤ 50%', 
             'Добавлен', 1.0, ''),
            ('17.11.2024', 'PETG', '750 г (1 кг), катушка Ø 100 мм', '1.75 мм', 'прозрачный', 
             'В герметичной упаковке, с влагопоглотителем, t 15–25°C, влажность ≤ 45%', 
             'Используется', 0.5, 'Текущий проект'),
            ('17.11.2024', 'NYLON', '750 г (1 кг), катушка Ø 100 мм', '1.75-3 мм', 'белый', 
             'В сухом месте, герметично, с силикагелем, t 15–25°C, влажность ≤ 30%', 
             'Израсходован', 0.0, 'Проект завершен')
        ]
        
        try:
            self.cursor.executemany('''
                INSERT INTO materials VALUES (NULL, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', sample_data)
            self.conn.commit()
        except Exception as e:
            print(f"Ошибка: {e}")

    def create_interface(self):
        main_frame = tk.Frame(self.root, bg='white')
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        header_frame = tk.Frame(main_frame, bg='#2B579A', height=40)
        header_frame.pack(fill=tk.X, pady=(0, 10))
        header_frame.pack_propagate(False)
        
        tk.Label(header_frame, text="📊 УЧЕТ МАТЕРИАЛОВ ДЛЯ 3D ПРИНТЕРОВ", 
                bg='#2B579A', fg='white', font=('Arial', 14, 'bold')).pack(pady=8)
        
        ribbon_frame = tk.Frame(main_frame, bg='#F2F2F2', height=50)
        ribbon_frame.pack(fill=tk.X, pady=(0, 10))
        ribbon_frame.pack_propagate(False)
        
        file_group = tk.Frame(ribbon_frame, bg='#F2F2F2')
        file_group.pack(side=tk.LEFT, padx=10)
        
        tk.Label(file_group, text="Файл:", bg='#F2F2F2', font=('Arial', 9, 'bold')).pack(side=tk.LEFT)
        
        tk.Button(file_group, text="Сохранить БД", command=self.save_database, 
                 bg='white', fg='black', font=('Arial', 8), relief='raised', bd=1,
                 width=12).pack(side=tk.LEFT, padx=2)
        tk.Button(file_group, text="Загрузить БД", command=self.load_database,
                 bg='white', fg='black', font=('Arial', 8), relief='raised', bd=1,
                 width=12).pack(side=tk.LEFT, padx=2)
        tk.Button(file_group, text="Новая БД", command=self.create_new_db,
                 bg='white', fg='black', font=('Arial', 8), relief='raised', bd=1,
                 width=10).pack(side=tk.LEFT, padx=2)
        
        data_group = tk.Frame(ribbon_frame, bg='#F2F2F2')
        data_group.pack(side=tk.LEFT, padx=20)
        
        tk.Label(data_group, text="Данные:", bg='#F2F2F2', font=('Arial', 9, 'bold')).pack(side=tk.LEFT)
        
        tk.Button(data_group, text="Добавить", command=self.add_item,
                 bg='#107C10', fg='white', font=('Arial', 8), relief='raised', bd=1,
                 width=10).pack(side=tk.LEFT, padx=2)
        tk.Button(data_group, text="Редактировать", command=self.edit_item,
                 bg='#2B579A', fg='white', font=('Arial', 8), relief='raised', bd=1,
                 width=12).pack(side=tk.LEFT, padx=2)
        tk.Button(data_group, text="Удалить", command=self.delete_item,
                 bg='#C00000', fg='white', font=('Arial', 8), relief='raised', bd=1,
                 width=8).pack(side=tk.LEFT, padx=2)
        
        export_group = tk.Frame(ribbon_frame, bg='#F2F2F2')
        export_group.pack(side=tk.LEFT, padx=20)
        
        tk.Label(export_group, text="Экспорт:", bg='#F2F2F2', font=('Arial', 9, 'bold')).pack(side=tk.LEFT)
        
        tk.Button(export_group, text="Excel", command=self.export_excel,
                 bg='#107C10', fg='white', font=('Arial', 8), relief='raised', bd=1,
                 width=8).pack(side=tk.LEFT, padx=2)
        tk.Button(export_group, text="Word", command=self.export_word,
                 bg='#2B579A', fg='white', font=('Arial', 8), relief='raised', bd=1,
                 width=8).pack(side=tk.LEFT, padx=2)
        tk.Button(export_group, text="Обновить", command=self.load_data,
                 bg='#666666', fg='white', font=('Arial', 8), relief='raised', bd=1,
                 width=10).pack(side=tk.LEFT, padx=2)
        
        filter_frame = tk.Frame(main_frame, bg='#E6E6E6', height=35)
        filter_frame.pack(fill=tk.X, pady=(0, 8))
        filter_frame.pack_propagate(False)
        
        filter_content = tk.Frame(filter_frame, bg='#E6E6E6')
        filter_content.pack(expand=True)
        
        tk.Label(filter_content, text="Фильтры:", bg='#E6E6E6', 
                font=('Arial', 9, 'bold')).grid(row=0, column=0, padx=5, pady=5, sticky='e')
        
        tk.Label(filter_content, text="Материал:", bg='#E6E6E6').grid(row=0, column=1, padx=2, pady=5, sticky='e')
        self.material_filter = ttk.Combobox(filter_content, values=['Все'], width=12, state='readonly')
        self.material_filter.grid(row=0, column=2, padx=2, pady=5, sticky='w')
        self.material_filter.bind('<<ComboboxSelected>>', self.apply_filters)
        
        tk.Label(filter_content, text="Сечение:", bg='#E6E6E6').grid(row=0, column=3, padx=5, pady=5, sticky='e')
        self.diameter_filter = ttk.Combobox(filter_content, values=['Все'], width=10, state='readonly')
        self.diameter_filter.grid(row=0, column=4, padx=2, pady=5, sticky='w')
        self.diameter_filter.bind('<<ComboboxSelected>>', self.apply_filters)
        
        tk.Label(filter_content, text="Цвет:", bg='#E6E6E6').grid(row=0, column=5, padx=5, pady=5, sticky='e')
        self.color_filter = ttk.Combobox(filter_content, values=['Все'], width=10, state='readonly')
        self.color_filter.grid(row=0, column=6, padx=2, pady=5, sticky='w')
        self.color_filter.bind('<<ComboboxSelected>>', self.apply_filters)
        
        tk.Label(filter_content, text="Статус:", bg='#E6E6E6').grid(row=0, column=7, padx=5, pady=5, sticky='e')
        self.status_filter = ttk.Combobox(filter_content, 
                                        values=['Все', 'Добавлен', 'Используется', 'Израсходован'], 
                                        width=12, state='readonly')
        self.status_filter.grid(row=0, column=8, padx=2, pady=5, sticky='w')
        self.status_filter.bind('<<ComboboxSelected>>', self.apply_filters)
        
        tk.Label(filter_content, text="Поиск:", bg='#E6E6E6').grid(row=0, column=9, padx=10, pady=5, sticky='e')
        self.search_var = tk.StringVar()
        search_entry = tk.Entry(filter_content, textvariable=self.search_var, width=20, 
                               font=('Arial', 9), relief='sunken', bd=1)
        search_entry.grid(row=0, column=10, padx=2, pady=5, sticky='w')
        search_entry.bind('<KeyRelease>', self.search_items)
        
        tk.Button(filter_content, text="Сбросить", command=self.reset_filters,
                 bg='#C00000', fg='white', font=('Arial', 8), relief='raised', bd=1,
                 width=8).grid(row=0, column=11, padx=10, pady=5)
        
        table_container = tk.Frame(main_frame, bg='#D0D0D0', relief='sunken', bd=1)
        table_container.pack(fill=tk.BOTH, expand=True)
        
        table_header = tk.Frame(table_container, bg='#2B579A', height=25)
        table_header.pack(fill=tk.X)
        table_header.pack_propagate(False)
        
        tk.Label(table_header, text="📋 ТАБЛИЦА МАТЕРИАЛОВ", 
                bg='#2B579A', fg='white', font=('Arial', 10, 'bold')).pack(pady=4)
        
        table_frame = tk.Frame(table_container, bg='white')
        table_frame.pack(fill=tk.BOTH, expand=True, padx=1, pady=1)
        
        columns = ("date", "material_type", "spool_size", "diameter", "color", 
                  "storage_conditions", "status", "remaining_weight", "notes")
        
        style = ttk.Style()
        style.theme_use('clam')
        style.configure("Treeview", 
                       background="white",
                       foreground="black",
                       rowheight=25,
                       fieldbackground="white",
                       font=('Arial', 9))
        style.configure("Treeview.Heading",
                       background="#2B579A",
                       foreground="white",
                       font=('Arial', 9, 'bold'),
                       relief='flat')
        style.map('Treeview.Heading', 
                 background=[('active', '#1E4A7F')])
        
        self.tree = ttk.Treeview(table_frame, columns=columns, show="headings", height=18)
        
        column_config = [
            ("date", "Дата", 90),
            ("material_type", "Вид материала", 120), 
            ("spool_size", "Размер катушки", 180),
            ("diameter", "Сечение", 80),
            ("color", "Цвет", 90),
            ("storage_conditions", "Условия хранения", 250),
            ("status", "Статус", 100),
            ("remaining_weight", "Остаток", 80),
            ("notes", "Примечания", 150)
        ]
        
        for col, header, width in column_config:
            self.tree.heading(col, text=header)
            self.tree.column(col, width=width, anchor='center', minwidth=width)
        
        v_scrollbar = ttk.Scrollbar(table_frame, orient=tk.VERTICAL, command=self.tree.yview)
        self.tree.configure(yscrollcommand=v_scrollbar.set)
        
        h_scrollbar = ttk.Scrollbar(table_frame, orient=tk.HORIZONTAL, command=self.tree.xview)
        self.tree.configure(xscrollcommand=h_scrollbar.set)
        
        self.tree.grid(row=0, column=0, sticky='nsew')
        v_scrollbar.grid(row=0, column=1, sticky='ns')
        h_scrollbar.grid(row=1, column=0, sticky='ew')
        
        table_frame.grid_rowconfigure(0, weight=1)
        table_frame.grid_columnconfigure(0, weight=1)
        
        self.tree.bind('<Double-1>', lambda e: self.edit_item())

    def search_items(self, event=None):
        if not hasattr(self, 'tree') or not self.tree:
            return
            
        search_text = self.search_var.get().lower()
        
        for item in self.tree.selection():
            self.tree.selection_remove(item)
        
        for item in self.tree.get_children():
            values = [str(v).lower() for v in self.tree.item(item)['values']]
            if any(search_text in value for value in values):
                self.tree.selection_add(item)
                self.tree.focus(item)

    def save_database(self):
        try:
            file_path = filedialog.asksaveasfilename(
                defaultextension=".db",
                filetypes=[("Database files", "*.db")],
                title="Сохранить базу данных"
            )
            
            if file_path:
                if self.conn:
                    self.conn.close()
                
                shutil.copy2(self.current_db, file_path)
                
                self.conn = sqlite3.connect(self.current_db)
                self.cursor = self.conn.cursor()
                
                messagebox.showinfo("Успех", f"База сохранена: {os.path.basename(file_path)}")
                
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка: {str(e)}")

    def load_database(self):
        try:
            file_path = filedialog.askopenfilename(
                filetypes=[("Database files", "*.db")],
                title="Выберите базу данных"
            )
            
            if file_path:
                if self.conn:
                    self.conn.close()
                
                self.current_db = file_path
                self.conn = sqlite3.connect(self.current_db)
                self.cursor = self.conn.cursor()
                
                self.load_data()
                
                messagebox.showinfo("Успех", f"База загружена: {os.path.basename(file_path)}")
                
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка: {str(e)}")

    def create_new_db(self):
        try:
            file_path = filedialog.asksaveasfilename(
                defaultextension=".db",
                filetypes=[("Database files", "*.db")],
                title="Создать новую базу"
            )
            
            if file_path:
                if self.conn:
                    self.conn.close()
                
                if os.path.exists(file_path):
                    os.remove(file_path)
                
                self.current_db = file_path
                self.conn = sqlite3.connect(self.current_db)
                self.cursor = self.conn.cursor()
                
                self.setup_database()
                
                messagebox.showinfo("Успех", f"Новая база создана: {os.path.basename(file_path)}")
                
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка: {str(e)}")

    def load_data(self):
        try:
            if not hasattr(self, 'tree') or not self.tree:
                return
                
            for item in self.tree.get_children():
                self.tree.delete(item)
            
            self.cursor.execute("SELECT * FROM materials ORDER BY date DESC")
            rows = self.cursor.fetchall()
            
            materials = set()
            diameters = set() 
            colors = set()
            
            for row in rows:
                materials.add(row[2])
                diameters.add(row[4])
                colors.add(row[5])
                
                remaining = row[8]
                if remaining >= 1:
                    remaining_text = f"{remaining:.0f} кг"
                else:
                    remaining_text = f"{remaining:.1f} кг"
                
                self.tree.insert("", tk.END, values=(
                    row[1], row[2], row[3], row[4], row[5], row[6], row[7], remaining_text, row[9]
                ))
            
            if hasattr(self, 'material_filter') and self.material_filter:
                self.material_filter['values'] = ['Все'] + sorted(materials)
            if hasattr(self, 'diameter_filter') and self.diameter_filter:
                self.diameter_filter['values'] = ['Все'] + sorted(diameters)
            if hasattr(self, 'color_filter') and self.color_filter:
                self.color_filter['values'] = ['Все'] + sorted(colors)
            
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка загрузки: {str(e)}")

    def apply_filters(self, event=None):
        try:
            if not hasattr(self, 'tree') or not self.tree:
                return
                
            query = "SELECT * FROM materials WHERE 1=1"
            params = []
            
            if hasattr(self, 'material_filter') and self.material_filter:
                material = self.material_filter.get()
                if material and material != 'Все':
                    query += " AND material_type = ?"
                    params.append(material)
            
            if hasattr(self, 'diameter_filter') and self.diameter_filter:
                diameter = self.diameter_filter.get()
                if diameter and diameter != 'Все':
                    query += " AND diameter = ?" 
                    params.append(diameter)
            
            if hasattr(self, 'color_filter') and self.color_filter:
                color = self.color_filter.get()
                if color and color != 'Все':
                    query += " AND color = ?"
                    params.append(color)
            
            if hasattr(self, 'status_filter') and self.status_filter:
                status = self.status_filter.get()
                if status and status != 'Все':
                    query += " AND status = ?"
                    params.append(status)
            
            self.cursor.execute(query, params)
            filtered_data = self.cursor.fetchall()
            
            for item in self.tree.get_children():
                self.tree.delete(item)
                
            for row in filtered_data:
                remaining = row[8]
                if remaining >= 1:
                    remaining_text = f"{remaining:.0f} кг"
                else:
                    remaining_text = f"{remaining:.1f} кг"
                
                self.tree.insert("", tk.END, values=(
                    row[1], row[2], row[3], row[4], row[5], row[6], row[7], remaining_text, row[9]
                ))
                
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка фильтрации: {str(e)}")

    def reset_filters(self):
        if hasattr(self, 'material_filter') and self.material_filter:
            self.material_filter.set('')
        if hasattr(self, 'diameter_filter') and self.diameter_filter:
            self.diameter_filter.set('')
        if hasattr(self, 'color_filter') and self.color_filter:
            self.color_filter.set('')
        if hasattr(self, 'status_filter') and self.status_filter:
            self.status_filter.set('')
        if hasattr(self, 'search_var') and self.search_var:
            self.search_var.set('')
        self.load_data()

    def add_item(self):
        self.show_material_dialog()

    def edit_item(self):
        if not hasattr(self, 'tree') or not self.tree:
            return
            
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("Внимание", "Выберите материал для редактирования")
            return
        
        item_data = self.tree.item(selected[0])['values']
        
        self.cursor.execute("SELECT * FROM materials WHERE date=? AND material_type=? AND color=?", 
                          (item_data[0], item_data[1], item_data[4]))
        db_data = self.cursor.fetchone()
        
        if db_data:
            self.show_material_dialog(db_data)

    def show_material_dialog(self, item_data=None):
        dialog = tk.Toplevel(self.root)
        dialog.title("Добавление материала" if not item_data else "Редактирование материала")
        dialog.geometry("450x550")
        dialog.configure(bg='white')
        dialog.resizable(False, False)
        dialog.transient(self.root)
        dialog.grab_set()
        
        header = tk.Frame(dialog, bg='#2B579A', height=40)
        header.pack(fill=tk.X)
        header.pack_propagate(False)
        
        title = "➕ ДОБАВЛЕНИЕ МАТЕРИАЛА" if not item_data else "✏️ РЕДАКТИРОВАНИЕ МАТЕРИАЛА"
        tk.Label(header, text=title, bg='#2B579A', fg='white', 
                font=('Arial', 12, 'bold')).pack(pady=10)
        
        form_frame = tk.Frame(dialog, bg='white')
        form_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)
        
        date_var = tk.StringVar(value=datetime.now().strftime('%d.%m.%Y'))
        material_var = tk.StringVar()
        spool_var = tk.StringVar()
        diameter_var = tk.StringVar()
        color_var = tk.StringVar()
        storage_var = tk.StringVar()
        status_var = tk.StringVar(value='Добавлен')
        remaining_var = tk.StringVar(value='1.0')
        notes_var = tk.StringVar()
        
        if item_data:
            date_var.set(item_data[0])
            material_var.set(item_data[1])
            spool_var.set(item_data[2])
            diameter_var.set(item_data[3])
            color_var.set(item_data[4])
            storage_var.set(item_data[5])
            status_var.set(item_data[6])
            remaining_var.set(item_data[7].replace(' кг', ''))
            notes_var.set(item_data[8])
        
        fields = [
            ("Дата:", date_var, 'entry'),
            ("Вид материала:", material_var, 'combo', ['ABS', 'PLA', 'PETG', 'NYLON', 'TPU', 'Другой']),
            ("Размер катушки:", spool_var, 'combo', ['750 г (1 кг), катушка Ø 100 мм', '1 кг, катушка Ø 100 мм', '2 кг, катушка Ø 150 мм']),
            ("Сечение:", diameter_var, 'combo', ['1.75 мм', '2.85 мм', '3 мм', '1.75-3 мм']),
            ("Цвет:", color_var, 'combo', ['чёрный', 'белый', 'прозрачный', 'синий', 'красный', 'зеленый', 'желтый', 'другой']),
            ("Условия хранения:", storage_var, 'entry'),
            ("Статус:", status_var, 'combo', ['Добавлен', 'Используется', 'Израсходован']),
            ("Остаток (кг):", remaining_var, 'entry'),
            ("Примечания:", notes_var, 'entry')
        ]
        
        for i, (label, var, field_type, *options) in enumerate(fields):
            tk.Label(form_frame, text=label, bg='white', font=('Arial', 9),
                    anchor='e').grid(row=i, column=0, sticky='e', pady=8, padx=5)
            
            if field_type == 'combo':
                entry = ttk.Combobox(form_frame, textvariable=var, values=options[0], 
                                   width=30, font=('Arial', 9))
            else:
                entry = tk.Entry(form_frame, textvariable=var, width=33, 
                               font=('Arial', 9), relief='sunken', bd=1)
            
            entry.grid(row=i, column=1, sticky='w', pady=8, padx=5)
        
        btn_frame = tk.Frame(form_frame, bg='white')
        btn_frame.grid(row=len(fields), column=0, columnspan=2, pady=20)
        
        tk.Button(btn_frame, text="Сохранить", command=lambda: self.save_material_dialog(
            date_var, material_var, spool_var, diameter_var, color_var, 
            storage_var, status_var, remaining_var, notes_var, item_data, dialog),
            bg='#107C10', fg='white', font=('Arial', 9), width=12, height=1).pack(side=tk.LEFT, padx=10)
        
        tk.Button(btn_frame, text="Отмена", command=dialog.destroy,
                 bg='#C00000', fg='white', font=('Arial', 9), width=12, height=1).pack(side=tk.LEFT, padx=10)

    def save_material_dialog(self, date_var, material_var, spool_var, diameter_var, color_var,
                           storage_var, status_var, remaining_var, notes_var, item_data, dialog):
        try:
            if not material_var.get():
                raise ValueError("Введите вид материала")
            
            remaining = float(remaining_var.get().replace(',', '.'))
            
            if item_data:
                self.cursor.execute('''
                    UPDATE materials 
                    SET date=?, material_type=?, spool_size=?, diameter=?, color=?,
                        storage_conditions=?, status=?, remaining_weight=?, notes=?
                    WHERE date=? AND material_type=? AND color=?
                ''', (date_var.get(), material_var.get(), spool_var.get(), diameter_var.get(),
                      color_var.get(), storage_var.get(), status_var.get(), remaining,
                      notes_var.get(), item_data[0], item_data[1], item_data[4]))
            else:
                self.cursor.execute('''
                    INSERT INTO materials 
                    VALUES (NULL, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                ''', (date_var.get(), material_var.get(), spool_var.get(), diameter_var.get(),
                      color_var.get(), storage_var.get(), status_var.get(), remaining,
                      notes_var.get()))
            
            self.conn.commit()
            self.load_data()
            dialog.destroy()
            messagebox.showinfo("Успех", "✅ Материал успешно сохранен")
            
        except Exception as e:
            messagebox.showerror("Ошибка", f"❌ Ошибка: {str(e)}")

    def delete_item(self):
        if not hasattr(self, 'tree') or not self.tree:
            return
            
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("Внимание", "Выберите материал для удаления")
            return
        
        item_data = self.tree.item(selected[0])['values']
        
        if messagebox.askyesno("Подтверждение", f"Удалить материал:\n{item_data[1]} - {item_data[4]}?"):
            try:
                self.cursor.execute("DELETE FROM materials WHERE date=? AND material_type=? AND color=?", 
                                  (item_data[0], item_data[1], item_data[4]))
                self.conn.commit()
                self.load_data()
                messagebox.showinfo("Успех", "✅ Материал удален")
            except Exception as e:
                messagebox.showerror("Ошибка", f"❌ Ошибка удаления: {str(e)}")

    def export_excel(self):
        try:
            file_path = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Excel files", "*.xlsx")],
                title="Экспорт в Excel"
            )
            
            if file_path:
                self.cursor.execute("SELECT * FROM materials")
                data = self.cursor.fetchall()
                
                excel_data = []
                for row in data:
                    remaining = row[8]
                    if remaining >= 1:
                        remaining_text = f"{remaining:.0f} кг"
                    else:
                        remaining_text = f"{remaining:.1f} кг"
                    
                    excel_data.append({
                        'Дата': row[1],
                        'Вид материала': row[2],
                        'Размер катушки, вес кг.': row[3],
                        'Сечение': row[4],
                        'Цвет': row[5],
                        'Условия хранения': row[6],
                        'Статус': row[7],
                        'Остаток': remaining_text
                    })
                
                df = pd.DataFrame(excel_data)
                df.to_excel(file_path, index=False)
                messagebox.showinfo("Успех", f"✅ Данные экспортированы в Excel")
                
        except Exception as e:
            messagebox.showerror("Ошибка", f"❌ Ошибка экспорта: {str(e)}")

    def export_word(self):
        try:
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word files", "*.docx")],
                title="Экспорт в Word"
            )
            
            if file_path:
                self.cursor.execute("SELECT * FROM materials")
                data = self.cursor.fetchall()
                
                doc = Document()
                doc.add_heading('Отчет по материалам для 3D печати', 0)
                doc.add_paragraph(f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
                doc.add_paragraph(f"Всего материалов: {len(data)}")
                
                table = doc.add_table(rows=1, cols=8)
                table.style = 'Table Grid'
                
                headers = ['Дата', 'Материал', 'Размер катушки', 'Сечение', 'Цвет', 
                          'Условия хранения', 'Статус', 'Остаток']
                
                for i, header in enumerate(headers):
                    table.rows[0].cells[i].text = header
                
                for row in data:
                    remaining = row[8]
                    if remaining >= 1:
                        remaining_text = f"{remaining:.0f} кг"
                    else:
                        remaining_text = f"{remaining:.1f} кг"
                    
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(row[1])
                    row_cells[1].text = str(row[2])
                    row_cells[2].text = str(row[3])
                    row_cells[3].text = str(row[4])
                    row_cells[4].text = str(row[5])
                    row_cells[5].text = str(row[6])
                    row_cells[6].text = str(row[7])
                    row_cells[7].text = remaining_text
                
                doc.save(file_path)
                messagebox.showinfo("Успех", f"✅ Отчет создан в Word")
                
        except Exception as e:
            messagebox.showerror("Ошибка", f"❌ Ошибка: {str(e)}")

def main():
    root = tk.Tk()
    app = FilamentApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()
